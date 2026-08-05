# ingest_hybrid.py
# ==========================================================
# ENTERPRISE PRODUCTION HYBRID INGESTION PIPELINE
# ==========================================================
# FEATURES
# ==========================================================
# âœ… PostgreSQL + pgvector storage
# âœ… Enterprise-grade ingestion
# âœ… Semantic deduplication
# âœ… Token-aware chunking
# âœ… BM25 hybrid retrieval support
# âœ… HNSW vector indexing
# âœ… Batch PostgreSQL inserts
# âœ… Stable deterministic IDs
# âœ… Metadata enrichment
# âœ… Recovery-safe ingestion
# âœ… Mobile-RAG optimized
# âœ… Large-scale production ready
# âœ… Fast retrieval optimized
# âœ… GPU embedding support
# âœ… ONNX embedding acceleration
# âœ… File hash deduplication
# âœ… Manifest tracking
# âœ… Retry-safe ingestion
# âœ… Chunk metadata
# âœ… Topic tagging
# âœ… Upload authentication
# âœ… Background queue processing
# âœ… PDF validation
# âœ… Failure recovery
# âœ… Version tracking
# âœ… Enterprise observability
# âœ… Async-safe ingestion
# âœ… Production-grade indexing
# ==========================================================

import os
import re
import gc
import json
import time
import hmac
import queue
import logging
import base64
import threading
import unicodedata
from concurrent.futures import ThreadPoolExecutor, as_completed

from pathlib import Path
from datetime import datetime

import docx
import fitz
import httpx
import psycopg2
import numpy as np
from simhash import Simhash

import requests

from Crypto.Hash import SHA256

from psycopg2.pool import SimpleConnectionPool
from psycopg2.extras import execute_batch

from rapidfuzz import fuzz

from langchain_text_splitters import RecursiveCharacterTextSplitter

# ==========================================================
# GEMINI API EMBEDDINGS
# ==========================================================

USE_ONNX = False

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
GEMINI_EMBED_URL = (
    "https://generativelanguage.googleapis.com/v1/"
    "models/gemini-embedding-001:batchEmbedContents"
)
GEMINI_EMBED_BATCH = int(
    os.getenv("GEMINI_EMBED_BATCH", "20")
)
GEMINI_EMBED_TASK = (
    os.getenv("GEMINI_EMBED_TASK", "RETRIEVAL_DOCUMENT")
)

# ==========================================================
# LOGGING
# ==========================================================

logging.basicConfig(
    level=os.getenv("LOG_LEVEL", "INFO"),
    format="%(asctime)s [%(levelname)s] %(message)s"
)

logger = logging.getLogger(__name__)

# ==========================================================
# CONFIG
# ==========================================================

APP_VERSION = "13.0.0"

UPLOAD_API_KEY = os.getenv(
    "UPLOAD_API_KEY",
    "change_this_key"
)

# ==========================================================
# EMBEDDING MODEL (Gemini API)
# ==========================================================

EMBED_PROVIDER = os.getenv("EMBED_PROVIDER", "gemini")

EMBED_MODEL = os.getenv(
    "EMBED_MODEL",
    "gemini-embedding-001"
)

EMBEDDING_RETRIES = int(
    os.getenv("EMBEDDING_RETRIES", "3")
)

EMBEDDING_RETRY_DELAY = float(
    os.getenv("EMBEDDING_RETRY_DELAY", "2")
)

# ==========================================================
# DEVICE
# ==========================================================

FORCE_CPU_ONLY = (
    os.getenv(
        "FORCE_CPU_ONLY",
        "false"
    ).lower() == "true"
)

DEVICE = "cpu"

logger.info(
    f"Embedding provider: Gemini API (gemini-embedding-001)"
)

# ==========================================================
# EMBEDDING DIM
# ==========================================================

EMBED_DIM = int(
    os.getenv("EMBED_DIM", "3072")
)

HEADING_PREFIX_MAX_LEVELS = int(
    os.getenv("HEADING_PREFIX_MAX_LEVELS", "2")
)

# ==========================================================
# CHUNKING
# ==========================================================
CHUNK_SIZE = int(
    os.getenv("CHUNK_SIZE", "1024")
)
CHUNK_OVERLAP = int(
    os.getenv("CHUNK_OVERLAP", "256")
)
# Fractional tolerance above the target chunk size before splitting is forced.
# Kept as a code constant (algorithmic behavior, not an operational setting).
CHUNK_SIZE_TOLERANCE = 0.15
CHUNK_SIZE_TOLERANT = int(
    round(CHUNK_SIZE * (1.0 + CHUNK_SIZE_TOLERANCE))
)

MIN_CHUNK_TOKENS = int(
    os.getenv("MIN_CHUNK_TOKENS", "8")
)

# ==========================================================
# PERFORMANCE
# ==========================================================

BATCH_SIZE = int(
    os.getenv("BATCH_SIZE", "512")
)

EMBED_CONCURRENCY = int(
    os.getenv("EMBED_CONCURRENCY", "4")
)

MAX_WORKERS = int(
    os.getenv("MAX_WORKERS", "4")
)

PDF_PAGE_BATCH_SIZE = int(
    os.getenv("PDF_PAGE_BATCH_SIZE", "50")
)

PDF_PARSE_TIMEOUT_SECONDS = int(
    os.getenv("PDF_PARSE_TIMEOUT_SECONDS", "120")
)

STREAM_BATCH_TOKENS = int(
    os.getenv("STREAM_BATCH_TOKENS", "2048")
)

DIAGRAM_ASSET_DIR = Path(
    os.getenv("DIAGRAM_ASSET_DIR", "/app/diagram-assets")
)
DIAGRAM_ASSET_URL_PREFIX = os.getenv(
    "DIAGRAM_ASSET_URL_PREFIX",
    "/diagram-assets"
).rstrip("/")
DIAGRAM_RENDER_SCALE = float(
    os.getenv("DIAGRAM_RENDER_SCALE", "1.4")
)
DIAGRAM_MIN_DRAWINGS = int(
    os.getenv("DIAGRAM_MIN_DRAWINGS", "3")
)
DIAGRAM_RENDER_MAX_PAGES = int(
    os.getenv("DIAGRAM_RENDER_MAX_PAGES", "120")
)

# ==========================================================
# VISUAL VERBALIZATION (Multimodal Image Description)
# ==========================================================

ENABLE_VISUAL_VERBALIZATION = (
    os.getenv("ENABLE_VISUAL_VERBALIZATION", "false").lower() == "true"
)
GEMINI_VERBALIZE_MODEL = os.getenv(
    "GEMINI_VERBALIZE_MODEL", "gemini-2.5-flash"
)
VERBALIZE_MAX_IMAGES_PER_PAGE = int(
    os.getenv("VERBALIZE_MAX_IMAGES_PER_PAGE", "3")
)
VERBALIZE_CONCURRENCY = int(
    os.getenv("VERBALIZE_CONCURRENCY", "2")
)
VERBALIZE_API_DELAY = float(
    os.getenv("VERBALIZE_API_DELAY", "3.0")
)
VERBALIZE_MIN_IMAGE_SIZE = int(
    os.getenv("VERBALIZE_MIN_IMAGE_SIZE", "5000")
)
ENABLE_TABLE_EXTRACTION = (
    os.getenv("ENABLE_TABLE_EXTRACTION", "true").lower() == "true"
)

MODEL_MAX_TOKENS = int(
    os.getenv("MODEL_MAX_TOKENS", "512")
)

DB_MAX_RETRIES = int(
    os.getenv("DB_MAX_RETRIES", "5")
)

DB_RETRY_DELAY = float(
    os.getenv("DB_RETRY_DELAY", "2")
)

POSTGRES_STATEMENT_TIMEOUT = int(
    os.getenv("POSTGRES_STATEMENT_TIMEOUT", "30000")
)

# ==========================================================
# POSTGRES
# ==========================================================

PG_DB = os.getenv(
    "DB_NAME",
    "aryabhata_db"
)

PG_USER = os.getenv(
    "DB_USER",
    "aryabhata_user"
)

PG_PASS = os.getenv(
    "DB_PASSWORD",
    "Password123"
)

PG_HOST = os.getenv(
    "DB_HOST",
    "postgres"
)

PG_PORT = os.getenv(
    "DB_PORT",
    "5432"
)

PG_POOL_MIN = int(
    os.getenv(
        "PG_POOL_MIN_CONN",
        "2"
    )
)

PG_POOL_MAX = int(
    os.getenv(
        "PG_POOL_MAX_CONN",
        "30"
    )
)

# ==========================================================
# DEDUP
# ==========================================================

_raw_dedup_threshold = float(
    os.getenv(
        "SIMILARITY_DEDUP_THRESHOLD",
        "92"
    )
)

SIMILARITY_DEDUP_THRESHOLD = (
    int(_raw_dedup_threshold * 100)
    if _raw_dedup_threshold <= 1
    else int(_raw_dedup_threshold)
)

SIMHASH_BUCKET_BITS = int(
    os.getenv("SIMHASH_BUCKET_BITS", "16")
)

DEDUP_COMPARE_LIMIT = int(
    os.getenv("DEDUP_COMPARE_LIMIT", "64")
)

CLEAR_CUDA_CACHE_AFTER_FILE = os.getenv(
    "CLEAR_CUDA_CACHE_AFTER_FILE",
    "false"
).lower() == "true"

# ==========================================================
# QUEUE
# ==========================================================

ingestion_queue = queue.Queue(
    maxsize=MAX_WORKERS * 4
)

embedding_lock = threading.Lock()

# ==========================================================
# GEMINI API EMBEDDING
# ==========================================================

_gemini_session = requests.Session()


def _gemini_embed_batch(texts, task_type="RETRIEVAL_DOCUMENT"):
    """Call Gemini batchEmbedContents API for a list of texts."""
    payload = {
        "requests": [
            {
                "model": f"models/{EMBED_MODEL}",
                "content": {"parts": [{"text": t}]},
                "taskType": task_type,
                "outputDimensionality": EMBED_DIM,
            }
            for t in texts
        ]
    }
    for attempt in range(8):
        try:
            resp = _gemini_session.post(
                f"{GEMINI_EMBED_URL}?key={GEMINI_API_KEY}",
                json=payload,
                timeout=90,
            )
            if resp.status_code == 429:
                logger.warning(
                    f"Gemini embed 429 body: {resp.text[:300]}"
                )
                wait = (2 ** attempt) * 2
                logger.warning(
                    f"Gemini embed 429, attempt {attempt+1}/8, retry in {wait}s"
                )
                time.sleep(wait)
                continue
            resp.raise_for_status()
            data = resp.json()
            return [
                e["values"] for e in data["embeddings"]
            ]
        except Exception as exc:
            if attempt == 7:
                raise
            wait = (2 ** attempt) * 2
            logger.warning(
                f"Gemini embed error ({exc}), attempt {attempt+1}/8, retry in {wait}s"
            )
            time.sleep(wait)
    raise RuntimeError("Failed to embed batch after 8 retries")

# ==========================================================
# POSTGRES POOL
# ==========================================================

logger.info(
    "Connecting PostgreSQL pool..."
)

pg_pool = SimpleConnectionPool(
    minconn=PG_POOL_MIN,
    maxconn=PG_POOL_MAX,
    dbname=PG_DB,
    user=PG_USER,
    password=PG_PASS,
    host=PG_HOST,
    port=PG_PORT
)

logger.info(
    "âœ… PostgreSQL pool ready"
)

# ==========================================================
# HELPERS
# ==========================================================

def get_conn():
    return pg_pool.getconn()

def release_conn(conn):
    pg_pool.putconn(conn)

def retry_operation(
    label,
    fn,
    retries,
    delay
):

    last_error = None

    for attempt in range(
        1,
        retries + 1
    ):

        try:

            return fn()

        except Exception as e:

            last_error = e

            logger.warning(
                f"{label} failed attempt {attempt}/{retries}: {e}"
            )

            if attempt < retries:
                is_429 = "429" in str(e)
                wait = delay * (2 ** (attempt - 1)) if is_429 else delay
                time.sleep(wait)

    raise last_error

# ==========================================================
# ENSURE TABLES
# ==========================================================

def ensure_tables():

    conn = get_conn()

    try:

        with conn.cursor() as cur:

            cur.execute(
                "CREATE EXTENSION IF NOT EXISTS vector;"
            )

            cur.execute(
                "CREATE EXTENSION IF NOT EXISTS pg_trgm;"
            )

            # ==================================================
            # DOCUMENTS
            # ==================================================

            cur.execute("""
                CREATE TABLE IF NOT EXISTS documents (
                    file_hash TEXT PRIMARY KEY,
                    file_name TEXT,
                    subject_id TEXT,
                    version INTEGER DEFAULT 1,
                    status TEXT,
                    error_message TEXT,
                    created_at TIMESTAMP DEFAULT NOW(),
                    updated_at TIMESTAMP DEFAULT NOW()
                );
            """)

            cur.execute("""
                ALTER TABLE documents
                ADD COLUMN IF NOT EXISTS error_message TEXT;
            """)

            cur.execute("""
                ALTER TABLE documents
                ADD COLUMN IF NOT EXISTS subject_id TEXT;
            """)

            # ==================================================
            # MANIFEST
            # ==================================================

            cur.execute("""
                CREATE TABLE IF NOT EXISTS ingestion_manifest (
                    file_hash TEXT PRIMARY KEY,
                    filename TEXT,
                    chunk_count INTEGER DEFAULT 0,
                    vector_inserted BOOLEAN DEFAULT FALSE,
                    postgres_inserted BOOLEAN DEFAULT FALSE,
                    ingestion_version INTEGER DEFAULT 1,
                    created_at TIMESTAMP DEFAULT NOW(),
                    updated_at TIMESTAMP DEFAULT NOW()
                );
            """)

            cur.execute("""
                ALTER TABLE ingestion_manifest
                ADD COLUMN IF NOT EXISTS vector_inserted BOOLEAN DEFAULT FALSE;
            """)

            cur.execute("""
                ALTER TABLE ingestion_manifest
                ADD COLUMN IF NOT EXISTS created_at TIMESTAMP DEFAULT NOW();
            """)

            cur.execute("""
                ALTER TABLE ingestion_manifest
                ADD COLUMN IF NOT EXISTS updated_at TIMESTAMP DEFAULT NOW();
            """)

            # ==================================================
            # CHUNKS — subject-partitioned table
            # ==================================================

            # Migrate old non-partitioned table
            cur.execute("""
                SELECT EXISTS (
                    SELECT 1 FROM pg_class
                    WHERE relname='upsc_chunks'
                    AND relkind='r'
                )
            """)

            old_table_is_regular = cur.fetchone()[0]

            if old_table_is_regular:

                cur.execute("""
                    ALTER TABLE IF EXISTS upsc_chunks
                    RENAME TO upsc_chunks_legacy
                """)

                logger.info(
                    "Migrated existing upsc_chunks to upsc_chunks_legacy"
                )

                cur.execute("""
                    DROP TABLE IF EXISTS upsc_chunks_legacy CASCADE
                """)

                # Drop and recreate partitioned table to ensure correct schema
                cur.execute("""
                    DROP TABLE IF EXISTS upsc_chunks CASCADE
                """)

            cur.execute(f"""
                CREATE TABLE IF NOT EXISTS upsc_chunks (
                    id TEXT NOT NULL,
                    chunk TEXT,
                    topic TEXT,
                    difficulty TEXT,
                    source_file TEXT,
                    file_hash TEXT,
                    subject_id TEXT NOT NULL DEFAULT 'general',
                    chunk_index INTEGER,
                    page_number INTEGER,

                    chunk_version INTEGER DEFAULT 1,
                    embedding VECTOR({EMBED_DIM}),
                    search_vector tsvector,
                    heading_hierarchy jsonb DEFAULT '[]'::jsonb,
                    parent_chunk TEXT DEFAULT '',
                    is_parent_chunk boolean DEFAULT false,
                    created_at TIMESTAMP DEFAULT NOW(),
                    PRIMARY KEY (id, subject_id)
                ) PARTITION BY LIST (subject_id);
            """)

            # ==================================================
            # CREATE PARTITIONS
            # ==================================================

            partition_subjects = [
                'polity', 'history', 'economy', 'geography',
                'environment', 'science', 'ethics', 'general'
            ]

            for subj in partition_subjects:

                cur.execute(f"""
                    CREATE TABLE IF NOT EXISTS upsc_chunks_{subj}
                    PARTITION OF upsc_chunks
                    FOR VALUES IN ('{subj}');
                """)

            # ==================================================
            # INDEXES on partitioned parent
            # ==================================================

            cur.execute("""
                CREATE INDEX IF NOT EXISTS idx_topic_partitioned
                ON upsc_chunks(topic);
            """)

            cur.execute("""
                CREATE INDEX IF NOT EXISTS idx_file_hash_partitioned
                ON upsc_chunks(file_hash);
            """)

            cur.execute("""
                CREATE INDEX IF NOT EXISTS idx_subject_id
                ON upsc_chunks(subject_id);
            """)

            cur.execute("""
                CREATE INDEX IF NOT EXISTS idx_gin_searchvec_partitioned
                ON upsc_chunks
                USING GIN(search_vector);
            """)

            cur.execute("""
                CREATE INDEX IF NOT EXISTS idx_gin_trgm_partitioned
                ON upsc_chunks
                USING GIN(chunk gin_trgm_ops);
            """)

            # ==================================================
            # HNSW INDEX
            # ==================================================

            if EMBED_DIM <= 2000:

                cur.execute(f"""
                    CREATE INDEX IF NOT EXISTS idx_hnsw_v_partitioned
                    ON upsc_chunks
                    USING hnsw (
                        embedding vector_cosine_ops
                    );
                """)

        conn.commit()

        logger.info(
            "âœ… PostgreSQL tables ensured"
        )

    except Exception as e:

        conn.rollback()

        logger.exception(
            f"âŒ ensure_tables failed: {e}"
        )

        raise

    finally:

        release_conn(conn)

# ==========================================================
# FILE HASH
# ==========================================================

def file_checksum(path: Path):

    h = SHA256.new()

    with open(path, "rb") as f:

        for chunk in iter(
            lambda: f.read(4096),
            b""
        ):
            h.update(chunk)

    return h.hexdigest()

# ==========================================================
# PDF VALIDATION
# ==========================================================

def validate_pdf(path: Path):

    try:

        with fitz.open(str(path)) as pdf:

            if pdf.page_count == 0:
                return False

        return True

    except:
        return False

# ==========================================================
# CLEAN TEXT
# ==========================================================

def clean_text(text: str):

    text = unicodedata.normalize(
        "NFKC",
        text
    )

    text = text.replace(
        "\u00a0",
        " "
    )

    # Preserve content inside ```text code blocks (ASCII diagrams) verbatim:
    # keep box-drawing characters, spacing, and alignment. Only clean the
    # text OUTSIDE code fences.
    segments = text.split("```")
    cleaned_segments = []

    for idx, segment in enumerate(segments):

        if idx % 2 == 1:

            # Inside a code fence — keep exactly as-is
            cleaned_segments.append(segment)
            continue

        # strip ASCII box-drawing characters
        segment = re.sub(
            r"[\u2500-\u257F\u2580-\u259F┌─│┐└┘├┤┬┴┼═║╒╓╔╕╖╗╘╙╚╛╜╝╞╟╠╡╢╣╤╥╦╧╨╩╪╫╬▀▁▄▅▆▇█▉▊▋▌▍▎▏■□▪▫]",
            "",
            segment,
        )

        # strip common watermark / header / footer lines
        segment = re.sub(
            r"(?im)^(page\s+\d+[\s\-_]*\d*|www\.\S+|\d+\s*/\s*\d+|for\s+(ias|upsc|prelims|mains).*?)$",
            "",
            segment,
        )

        segment = re.sub(
            r"\s+",
            " ",
            segment
        )

        cleaned_segments.append(segment)

    return "```".join(cleaned_segments).strip()

# ==========================================================
# READERS
# ==========================================================

def read_txt(path: Path):
    text = path.read_text(encoding="utf-8", errors="ignore")
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\u00a0", " ")
    text = re.sub(r"\r\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"Page\s+\d+", "", text, flags=re.IGNORECASE)
    return text.strip()

def read_txt_blocks(path: Path):
    block_lines = []
    block_tokens = 0
    block_index = 0

    with open(str(path), "r", encoding="utf-8", errors="ignore") as f:
        for line in f:
            block_lines.append(line)
            block_tokens += max(1, len(line.split()))

            if block_tokens >= STREAM_BATCH_TOKENS:
                text = "".join(block_lines)
                text = unicodedata.normalize("NFKC", text)
                text = text.replace("\u00a0", " ")
                text = re.sub(r"\r\n", "\n", text)
                text = re.sub(r"Page\s+\d+", "", text, flags=re.IGNORECASE)
                text = text.strip()
                if text:
                    yield {
                        "block_index": block_index,
                        "text": text
                    }
                block_lines.clear()
                block_tokens = 0
                block_index += 1

    if block_lines:
        text = "".join(block_lines)
        text = unicodedata.normalize("NFKC", text)
        text = text.replace("\u00a0", " ")
        text = re.sub(r"\r\n", "\n", text)
        text = re.sub(r"Page\s+\d+", "", text, flags=re.IGNORECASE)
        text = text.strip()
        if text:
            yield {
                "block_index": block_index,
                "text": text
            }

def read_docx(path: Path):
    document = docx.Document(str(path))
    texts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    return "\n\n".join(texts)

def read_docx_blocks(path: Path):
    document = docx.Document(str(path))
    block_paras = []
    block_tokens = 0
    block_index = 0

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        block_paras.append(text)
        block_tokens += max(1, len(text.split()))

        if block_tokens >= STREAM_BATCH_TOKENS:
            yield {
                "block_index": block_index,
                "text": "\n\n".join(block_paras)
            }
            block_paras.clear()
            block_tokens = 0
            block_index += 1

    if block_paras:
        yield {
            "block_index": block_index,
            "text": "\n\n".join(block_paras)
        }

def read_pdf(path: Path):
    if not validate_pdf(path):
        raise Exception("Invalid PDF file")
    pages = []
    with fitz.open(str(path)) as pdf:
        for page in pdf:
            txt = page.get_text().strip()
            if txt:
                pages.append(txt)
    text = "\n\n".join(pages)
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\u00a0", " ")
    text = re.sub(r"Page\s+\d+", "", text, flags=re.IGNORECASE)
    return text.strip()

MAX_LAYOUT_PAGES = int(os.getenv("MAX_LAYOUT_PAGES", "40"))
MAX_LAYOUT_FILE_MB = float(os.getenv("MAX_LAYOUT_FILE_MB", "10.0"))
USE_PYMUPDF4LLM = os.getenv("USE_PYMUPDF4LLM", "false").lower() == "true"


def read_pdf_pages(path: Path):
    if not validate_pdf(path):
        raise Exception("Invalid PDF file")

    file_size_mb = path.stat().st_size / (1024 * 1024)
    doc = fitz.open(str(path))
    total_pages = len(doc)

    logger.info(
        f"Processing PDF: {path.name} "
        f"({total_pages} pages, {file_size_mb:.2f} MB)"
    )

    if (
        USE_PYMUPDF4LLM
        and total_pages <= MAX_LAYOUT_PAGES
        and file_size_mb <= MAX_LAYOUT_FILE_MB
    ):
        try:
            import pymupdf4llm
            logger.info(
                f"Using pymupdf4llm (layout extraction) for {path.name}"
            )
            page_data = pymupdf4llm.to_markdown(
                doc,
                header=False,
                footer=False,
                show_progress=False,
                page_chunks=True
            )

            for page_record in page_data:
                page_text = page_record.get("text", "").strip()
                if not page_text:
                    continue

                page_number = page_record.get("page", 1)

                page_text = page_text.replace("\x00", "")
                page_text = unicodedata.normalize(
                    "NFKC", page_text
                )
                page_text = page_text.replace("\u00a0", " ")
                page_text = re.sub(
                    r"Page\s+\d+", "", page_text,
                    flags=re.IGNORECASE
                )
                page_text = page_text.strip()

                if page_text:
                    yield {
                        "page_number": page_number,
                        "text": page_text
                    }

            doc.close()
            return

        except MemoryError:
            logger.warning(
                "pymupdf4llm memory error, falling back to PyMuPDF"
            )
        except Exception as e:
            logger.warning(
                f"pymupdf4llm failed: {e}, "
                f"falling back to PyMuPDF"
            )

    logger.info(
        f"Using PyMuPDF blocks-stream (low memory) for {path.name}"
    )

    for page_index in range(total_pages):
        page = doc.load_page(page_index)
        blocks = page.get_text("blocks")

        page_lines = []
        for b in blocks:
            if b[6] == 0:
                clean_block = b[4].strip()
                if clean_block:
                    page_lines.append(clean_block)

        # Extract images and tables before nulling the page object
        images_data = []
        if ENABLE_VISUAL_VERBALIZATION:
            images_data = extract_page_images_base64(
                page, VERBALIZE_MAX_IMAGES_PER_PAGE, VERBALIZE_MIN_IMAGE_SIZE
            )

        tables_md = []
        if ENABLE_TABLE_EXTRACTION:
            tables_md = extract_tables_markdown(page)

        page = None

        if page_lines:
            page_content = "\n\n".join(page_lines)
            page_content = page_content.replace("\x00", "")
            page_content = unicodedata.normalize(
                "NFKC", page_content
            )
            page_content = page_content.replace("\u00a0", " ")
            page_content = re.sub(
                r"Page\s+\d+", "", page_content,
                flags=re.IGNORECASE
            )
            page_content = page_content.strip()

            # Inject Markdown tables into page text
            if tables_md:
                page_content += "\n\n" + "\n\n".join(tables_md)

            if page_content:
                yield {
                    "page_number": page_index + 1,
                    "text": page_content,
                    "images": images_data if ENABLE_VISUAL_VERBALIZATION else [],
                }

        if (page_index + 1) % 20 == 0:
            gc.collect()

    doc.close()
    gc.collect()

def page_has_visual_content(page):
    try:
        if page.get_images(full=True):
            return True
        return len(page.get_drawings()) >= DIAGRAM_MIN_DRAWINGS
    except Exception:
        return False

def render_pdf_visual_pages(path: Path, file_hash: str):
    page_urls = {}
    asset_dir = DIAGRAM_ASSET_DIR / file_hash
    asset_dir.mkdir(parents=True, exist_ok=True)

    with fitz.open(str(path)) as pdf:
        page_limit = min(pdf.page_count, DIAGRAM_RENDER_MAX_PAGES)
        matrix = fitz.Matrix(DIAGRAM_RENDER_SCALE, DIAGRAM_RENDER_SCALE)

        for page_index in range(page_limit):
            page = pdf[page_index]
            page_number = page_index + 1
            if not page_has_visual_content(page):
                continue

            filename = f"page-{page_number:04d}.png"
            output_path = asset_dir / filename
            if not output_path.exists():
                pix = page.get_pixmap(matrix=matrix, alpha=False)
                pix.save(str(output_path))

            page_urls[page_number] = (
                f"{DIAGRAM_ASSET_URL_PREFIX}/{file_hash}/{filename}"
            )

    return page_urls

# ==========================================================
# MULTIMODAL VISUAL EXTRACTION
# ==========================================================

def extract_page_images_base64(page, max_images=3, min_size_bytes=5000):
    """Extract meaningful images from a PDF page as base64 for Gemini API.
    Filters out small images (icons, logos, decorations) by byte size.
    """
    images = []
    try:
        for img_info in page.get_images(full=True):
            if len(images) >= max_images:
                break
            xref = img_info[0]
            base_image = page.parent.extract_image(xref)
            img_bytes = base_image["image"]
            if len(img_bytes) < min_size_bytes:
                continue
            ext = base_image.get("ext", "png")
            b64_data = base64.b64encode(img_bytes).decode("utf-8")
            images.append({
                "data": b64_data,
                "mime": f"image/{ext}",
                "size": len(img_bytes),
            })
    except Exception as e:
        logger.warning(f"Failed to extract images from page: {e}")
    return images


def extract_tables_markdown(page):
    """Extract tables from a PDF page as Markdown strings.
    Filters out text-box misdetections (single-column), header-only tables,
    and empty-header tables that PyMuPDF's find_tables() sometimes returns.
    """
    tables = []
    try:
        for table in page.find_tables():
            data = table.extract()
            if not data or len(data) < 2:
                continue

            # Normalize cells: None -> ""
            data = [[str(c or "") for c in row] for row in data]

            # Drop fully-empty leading/trailing columns so the table is tight.
            ncols = max(len(row) for row in data)
            data = [row + [""] * (ncols - len(row)) for row in data]
            col_has_content = [
                any(row[c].strip() for row in data)
                for c in range(ncols)
            ]
            if not any(col_has_content):
                continue
            first = col_has_content.index(True)
            last = len(col_has_content) - 1 - col_has_content[::-1].index(True)
            data = [[row[c] for c in range(first, last + 1)] for row in data]

            # Require a real multi-column table; single-column "tables" are
            # usually side-note text boxes, not tabular data.
            if len(data[0]) < 2:
                continue

            # Require a non-empty header row.
            if not any(data[0][c].strip() for c in range(len(data[0]))):
                continue

            header = data[0]
            md_rows = []
            md_rows.append("| " + " | ".join(header) + " |")
            md_rows.append("| " + " | ".join("---" for _ in header) + " |")
            for row in data[1:]:
                md_rows.append("| " + " | ".join(row) + " |")
            tables.append("\n".join(md_rows))
    except Exception as e:
        logger.warning(f"Failed to extract tables: {e}")
    return tables


# ==========================================================
# GEMINI VISUAL VERBALIZATION
# ==========================================================

_verbalize_session = requests.Session()


def verbalize_page_images_batch(pages_with_images):
    """Call Gemini Flash to describe all images on each page.
    Sends ONE API call per page (all images batched in a single request)
    to minimize quota consumption. Returns dict of page_number -> description.
    """
    if not pages_with_images:
        return {}

    model = GEMINI_VERBALIZE_MODEL
    url = (
        f"https://generativelanguage.googleapis.com/v1/"
        f"models/{model}:generateContent"
    )

    results = {}
    total_images = sum(len(item["images"]) for item in pages_with_images)

    logger.info(
        f"Verbalizing {total_images} images across "
        f"{len(pages_with_images)} pages (1 API call per page)..."
    )

    for idx, item in enumerate(pages_with_images):
        page_num = item["page_number"]
        images = item["images"]
        context = item.get("context", "")[:600]

        # Build one prompt with ALL images for this page in one API call
        parts = []
        prompt_text = (
            "You are part of a RAG ingestion pipeline. Examine ALL images below "
            "from a single page of an exam-preparation PDF. "
            "For each image, provide a concise description focusing on: "
            "text labels, data values, relationships, and key takeaways. "
            "Prefix each with 'Image N:'.\n\n"
            f"Page context:\n{context}\n"
            f"Total images on this page: {len(images)}"
        )
        parts.append({"text": prompt_text})

        for img in images:
            parts.append({
                "inline_data": {
                    "mime_type": img["mime"],
                    "data": img["data"],
                }
            })

        payload = {
            "contents": [{"parts": parts}],
            "generationConfig": {
                "temperature": 0.2,
                "maxOutputTokens": 768,
            }
        }

        desc = None
        for attempt in range(3):
            try:
                resp = _verbalize_session.post(
                    f"{url}?key={GEMINI_API_KEY}",
                    json=payload,
                    timeout=90,
                )
                if resp.status_code == 429:
                    wait = (2 ** attempt) * 4
                    logger.warning(
                        f"Verbalize 429 page {page_num}, "
                        f"attempt {attempt+1}/3, retry in {wait}s"
                    )
                    time.sleep(wait)
                    continue
                resp.raise_for_status()
                data = resp.json()
                text = (
                    data.get("candidates", [{}])[0]
                    .get("content", {})
                    .get("parts", [{}])[0]
                    .get("text", "")
                )
                if text:
                    desc = text.strip()
                break
            except Exception as e:
                if attempt == 2:
                    logger.warning(
                        f"Verbalization failed page {page_num}: {e}"
                    )
                else:
                    wait = (2 ** attempt) * 4
                    time.sleep(wait)

        if desc:
            results[page_num] = desc

        # Delay between API calls to avoid rate limiting
        if idx < len(pages_with_images) - 1:
            time.sleep(VERBALIZE_API_DELAY)

    return results


def enhance_batch_with_visuals(page_records, texts_list):
    """Verbalize images in page records and inject descriptions into text.
    Groups all images per page into a single API call for quota efficiency.
    """
    pages_to_verbalize = []
    for pr in page_records:
        pr_imgs = pr.get("images") or []
        if pr_imgs:
            pages_to_verbalize.append({
                "page_number": pr["page_number"],
                "images": pr_imgs,
                "context": (pr["text"] or "")[:600],
            })

    if not pages_to_verbalize:
        return

    verbalized = verbalize_page_images_batch(pages_to_verbalize)

    if not verbalized:
        logger.info("No verbalizations returned from Gemini")
    else:
        injected = 0
        for i, pr in enumerate(page_records):
            pn = pr["page_number"]
            desc = verbalized.get(pn)
            if desc:
                enhanced = pr["text"] + "\n\n📊 " + desc
                pr["text"] = enhanced
                texts_list[i] = enhanced
                injected += 1
        logger.info(f"Injected {injected} verbalized image descriptions")

    # Clean up image data to free memory
    for pr in page_records:
        pr.pop("images", None)

    gc.collect()


def infer_chunk_page_number(chunk, page_records):
    if not page_records:
        return None

    chunk_words = {
        word
        for word in normalized_words(chunk)
        if len(word) > 3
    }

    if not chunk_words:
        return None

    best_page = None
    best_score = 0

    for record in page_records:
        page_words = {
            word
            for word in normalized_words(record["text"])
            if len(word) > 3
        }
        score = len(chunk_words.intersection(page_words))
        if score > best_score:
            best_score = score
            best_page = record["page_number"]

    return best_page if best_score >= 5 else None

# ==========================================================
# TOKEN CHUNKING
# ==========================================================

def token_count(text):

    return max(
        1,
        len(
            re.findall(
                r"\w+|[^\w\s]",
                text,
                flags=re.UNICODE
            )
        )
    )

def token_limited_overlap(sentences):

    overlap = []

    total_tokens = 0

    for sentence in reversed(sentences):

        sentence_tokens = token_count(sentence)

        if total_tokens + sentence_tokens > CHUNK_OVERLAP:
            break

        overlap.append(sentence)

        total_tokens += sentence_tokens

    overlap.reverse()

    return overlap, total_tokens

def normalized_words(text):

    return re.findall(
        r"[a-z0-9]+",
        text.lower()
    )

def chunk_simhash(text):
    return Simhash(text).value

def simhash_bucket(fingerprint):

    if SIMHASH_BUCKET_BITS <= 0:
        return 0

    return fingerprint >> (
        64 - min(
            SIMHASH_BUCKET_BITS,
            64
        )
    )

def clean_ascii_diagrams(text):
    """Preserve ASCII/Unicode box diagrams as Markdown ```text blocks so chunking
    keeps them intact and LLMs recognize them as structured diagrams, not noise."""
    BOX_DRAWING_RE = re.compile(r"[\u2500-\u257F]")
    DIAGRAM_ARROW_RE = re.compile(r"[\u2190-\u21FF\u25B2\u25B3\u25BC\u25BD\u25B6\u25B7\u25C0\u25C1\u2B06\u2B07\u2B05\u2B08]")
    # Arrow used as a bullet/list marker (e.g. "⇒ Iron Nails, ...") is NOT a diagram.
    ARROW_BULLET_RE = re.compile(
        r"^\s*[\u2190-\u21FF\u25B2\u25B3\u25BC\u25BD\u25B6\u25B7\u25C0\u25C1\u2B06\u2B07\u2B05\u2B08]+\s+\S"
    )

    # If the document has no actual box-drawing characters anywhere, arrow lines
    # are almost certainly bullets/flow text, not diagrams — do not fence them.
    has_any_box = bool(BOX_DRAWING_RE.search(text))

    def is_diagram_line(line):
        if BOX_DRAWING_RE.search(line):
            return True
        if not has_any_box:
            return False
        if ARROW_BULLET_RE.match(line):
            return False
        return bool(DIAGRAM_ARROW_RE.search(line))

    lines = text.split("\n")
    has_box = [is_diagram_line(line) for line in lines]

    in_diagram = [False] * len(lines)
    for i in range(len(lines)):
        if has_box[i]:
            in_diagram[i] = True
            continue
        if not lines[i].strip():
            continue
        prev = i - 1
        while prev >= 0 and not lines[prev].strip():
            prev -= 1
        nxt = i + 1
        while nxt < len(lines) and not lines[nxt].strip():
            nxt += 1
        if prev >= 0 and nxt < len(lines) and has_box[prev] and has_box[nxt] and len(lines[i].strip()) < 300:
            in_diagram[i] = True

    result = []
    i = 0
    n = len(lines)
    while i < n:
        if in_diagram[i]:
            block = []
            while i < n and in_diagram[i]:
                block.append(lines[i])
                i += 1
            while block and not block[0].strip():
                block.pop(0)
            while block and not block[-1].strip():
                block.pop()
            if len(block) >= 2:
                result.append("```text")
                result.extend(block)
                result.append("```")
            else:
                result.extend(block)
            continue
        result.append(lines[i])
        i += 1

    return "\n".join(result)


def _split_parent_heading_aware(parent_text, parent_hierarchy):
    """Split an oversized parent into heading-delimited sub-segments.

    Returns a list of (hierarchy, text) tuples. Each sub-segment's hierarchy
    is the full ancestral chain (parent ancestors + the sub-headings it
    contains), deduplicated against the parent chain.
    """
    lines = parent_text.split("\n")
    heading_stack = []  # [(level, heading_text), ...]
    segments = []  # [(hierarchy, text), ...]
    current_lines = []
    parent_hier = list(parent_hierarchy or [])

    def flush():
        if not current_lines:
            return
        hier = list(parent_hier)
        for _, htext in heading_stack:
            if hier and hier[-1] == htext:
                continue
            hier.append(htext)
        seg_text = "\n".join(current_lines).strip()
        if seg_text:
            segments.append((hier, seg_text))
        current_lines.clear()

    for line in lines:
        stripped = line.strip()
        m = re.match(r"^(#{1,6})\s+(.+)", stripped)
        if m:
            flush()
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            while heading_stack and heading_stack[-1][0] >= level:
                heading_stack.pop()
            heading_stack.append((level, heading_text))
        current_lines.append(line)

    flush()

    if not segments:
        return [(list(parent_hier), parent_text)]
    return segments


def chunk_text(text):
    # ======================================================
    # GEMINI RAG CHUNKING — Layout-Aware with Ancestral
    # Heading Tracking and Parent-Child Design Pattern.
    #
    # Parent chunks: Full section text (~1024 tokens) for Gemini generation
    # Child chunks:  Smaller indexed blocks for vector/BM25 search
    # Each chunk carries its heading hierarchy as metadata.
    # ======================================================

    text = clean_ascii_diagrams(text)

    # --- Phase 1: Parse document into heading-aware sections ---
    lines = text.split("\n")
    heading_stack = []  # [(level, heading_text), ...]
    sections = []  # [(heading_hierarchy, section_text), ...]
    current_section_text = []
    current_hierarchy = []

    for line in lines:
        stripped = line.strip()
        heading_match = re.match(r"^(#{1,6})\s+(.+)", stripped)
        if heading_match:
            if current_section_text:
                sections.append((list(current_hierarchy), "\n".join(current_section_text).strip()))
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()

            embedded = re.findall(r"(?<=\S)\s*#{1,6}\s+(.+)", heading_text)
            if embedded:
                clean_heading = re.sub(r"\s*#{1,6}\s+.+", "", heading_text).strip()
                stripped = f"{'#' * level} {clean_heading}"
                heading_text = clean_heading

            while heading_stack and heading_stack[-1][0] >= level:
                heading_stack.pop()
            heading_stack.append((level, heading_text))

            if embedded:
                for extra in embedded:
                    clean_extra = extra.strip()
                    heading_stack.append((level, clean_extra))

            current_hierarchy = [h[1] for h in heading_stack]
            current_section_text = [stripped]
        else:
            if stripped or current_section_text:
                current_section_text.append(line)

    if current_section_text:
        sections.append((list(current_hierarchy), "\n".join(current_section_text).strip()))

    if not sections:
        return []

    # --- Phase 2: Build parent chunks from sections ---
    parent_chunks = []
    for hierarchy, section_text in sections:
        if not section_text:
            continue
        tk = token_count(section_text)
        if tk < MIN_CHUNK_TOKENS:
            continue
        parent_chunks.append({
            "text": section_text,
            "tokens": tk,
            "heading_hierarchy": hierarchy,
        })

    # Merge small consecutive sections to form properly-sized parent chunks
    merged_parents = []
    buf_text = []
    buf_tokens = 0
    buf_hierarchy = []

    for pc in parent_chunks:
        if buf_tokens + pc["tokens"] <= CHUNK_SIZE_TOLERANT:
            buf_text.append(pc["text"])
            buf_tokens += pc["tokens"]
            if not buf_hierarchy:
                buf_hierarchy = pc["heading_hierarchy"]
        else:
            if buf_text:
                merged_parents.append({
                    "text": "\n\n".join(buf_text),
                    "tokens": buf_tokens,
                    "heading_hierarchy": buf_hierarchy,
                })
            buf_text = [pc["text"]]
            buf_tokens = pc["tokens"]
            buf_hierarchy = pc["heading_hierarchy"]

    if buf_text:
        merged_parents.append({
            "text": "\n\n".join(buf_text),
            "tokens": buf_tokens,
            "heading_hierarchy": buf_hierarchy,
        })

    # --- Phase 3: Create child chunks from parent chunks ---
    results = []
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=token_count,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    for parent in merged_parents:
        parent_text = parent["text"]
        parent_hierarchy = parent["heading_hierarchy"]
        parent_tk = parent["tokens"]

        # If parent fits (within tolerance), it's both parent and child
        if parent_tk <= CHUNK_SIZE_TOLERANT:
            prefixed = _prefix_headings(parent_hierarchy, parent_text)
            results.append({
                "chunk_text": prefixed,
                "parent_text": parent_text,
                "heading_hierarchy": parent_hierarchy,
                "is_parent_chunk": True,
            })
            continue

        # Split parent into heading-delimited sub-segments (### -> ## -> #)
        heading_segments = _split_parent_heading_aware(
            parent_text,
            parent_hierarchy
        )

        # Greedily merge adjacent small sub-segments up to the tolerant bound
        merged_segments = []
        buf_hierarchy = []
        buf_text = []
        buf_tokens = 0

        for seg_hierarchy, seg_text in heading_segments:
            seg_tokens = token_count(seg_text)
            if buf_tokens + seg_tokens <= CHUNK_SIZE_TOLERANT:
                if not buf_hierarchy:
                    buf_hierarchy = seg_hierarchy
                buf_text.append(seg_text)
                buf_tokens += seg_tokens
            else:
                if buf_text:
                    merged_segments.append((buf_hierarchy, "\n\n".join(buf_text)))
                buf_hierarchy = seg_hierarchy
                buf_text = [seg_text]
                buf_tokens = seg_tokens

        if buf_text:
            merged_segments.append((buf_hierarchy, "\n\n".join(buf_text)))

        # Each merged segment becomes a child chunk; only a segment that still
        # exceeds the tolerant bound falls back to paragraph/sentence/token split.
        for seg_hierarchy, seg_text in merged_segments:
            seg_tokens = token_count(seg_text)
            if seg_tokens <= CHUNK_SIZE_TOLERANT:
                pieces = [seg_text]
            else:
                pieces = splitter.split_text(seg_text)
            for piece in pieces:
                piece = piece.strip()
                if not piece:
                    continue
                prefixed = _prefix_headings(seg_hierarchy, piece)
                results.append({
                    "chunk_text": prefixed,
                    "parent_text": parent_text,
                    "heading_hierarchy": seg_hierarchy,
                    "is_parent_chunk": False,
                })

    # ======================================================
    # DEDUP
    # ======================================================

    final_chunks = []
    dedup_buckets = {}
    exact_seen = set()

    for item in results:
        chunk = clean_text(item["chunk_text"])
        if len(chunk.split()) < 8:
            continue

        exact_key = " ".join(normalized_words(chunk))
        if exact_key in exact_seen:
            continue

        fingerprint = chunk_simhash(chunk)
        bucket = simhash_bucket(fingerprint)
        candidates = dedup_buckets.get(bucket, [])[-DEDUP_COMPARE_LIMIT:]

        duplicate = False
        for existing in candidates:
            score = fuzz.ratio(chunk, existing["chunk_text"])
            if score > SIMILARITY_DEDUP_THRESHOLD:
                duplicate = True
                break

        if duplicate:
            continue

        exact_seen.add(exact_key)
        item["chunk_text"] = chunk
        final_chunks.append(item)
        dedup_buckets.setdefault(bucket, []).append(item)

    return final_chunks


def _prefix_headings(hierarchy, text):
    """Prepend ancestral heading context to chunk text for Gemini."""
    if not hierarchy:
        return text
    clean_parts = [part.replace("#", "").strip() for part in hierarchy]
    if HEADING_PREFIX_MAX_LEVELS > 0:
        clean_parts = clean_parts[-HEADING_PREFIX_MAX_LEVELS:]
    prefix = " > ".join(clean_parts)
    return f"[Context: {prefix}]\n\n{text}"

# ==========================================================
# TOPIC TAGGING
# ==========================================================

def detect_topic(text):

    t = text.lower()

    if any(
        k in t
        for k in [
            "constitution",
            "parliament",
            "fundamental rights"
        ]
    ):
        return "Polity"

    if any(
        k in t
        for k in [
            "budget",
            "inflation",
            "gdp"
        ]
    ):
        return "Economy"

    if any(
        k in t
        for k in [
            "climate",
            "pollution",
            "forest"
        ]
    ):
        return "Environment"

    if any(
        k in t
        for k in [
            "freedom struggle",
            "mughal",
            "british india"
        ]
    ):
        return "History"

    if any(
        k in t
        for k in [
            "monsoon",
            "river",
            "mountain"
        ]
    ):
        return "Geography"

    return "General"

# ==========================================================
# DIFFICULTY
# ==========================================================

def detect_difficulty(text):

    words = len(
        text.split()
    )

    if words < 50:
        return "easy"

    if words < 150:
        return "medium"

    return "hard"

# ==========================================================
# AUTH
# ==========================================================

def validate_upload_auth():

    incoming = os.getenv(
        "UPLOAD_SECRET"
    )

    if not hmac.compare_digest(
        incoming or "",
        UPLOAD_API_KEY
    ):

        raise Exception(
            "Unauthorized upload"
        )

# ==========================================================
# DOCUMENT STATUS
# ==========================================================

def update_document_status(
    file_hash,
    file_name,
    status,
    error_message=None
):

    conn = get_conn()

    try:

        with conn.cursor() as cur:

            cur.execute("""
                INSERT INTO documents (
                    file_hash,
                    file_name,
                    status,
                    error_message
                )
                VALUES (%s, %s, %s, %s)

                ON CONFLICT (file_hash)

                DO UPDATE SET
                    status=EXCLUDED.status,
                    error_message=EXCLUDED.error_message,
                    updated_at=NOW()
            """, (
                file_hash,
                file_name,
                status,
                error_message
            ))

        conn.commit()

    finally:

        release_conn(conn)


def set_document_subject_id(
    file_hash,
    subject_id
):

    conn = get_conn()

    try:

        with conn.cursor() as cur:

            cur.execute("""
                UPDATE documents
                SET subject_id=%s
                WHERE file_hash=%s
            """, (
                subject_id,
                file_hash
            ))

        conn.commit()

    finally:

        release_conn(conn)

# ==========================================================
# DOCUMENT CHECK
# ==========================================================

def is_document_indexed(file_hash):

    conn = get_conn()

    try:

        with conn.cursor() as cur:

            cur.execute(
                "SELECT 1 FROM documents WHERE file_hash=%s AND status='indexed'",
                (file_hash,)
            )

            return cur.fetchone() is not None

    finally:

        release_conn(conn)

# ==========================================================
# DELETE EXISTING
# ==========================================================

def delete_existing_file_data(
    file_hash
):

    logger.info(
        f"ðŸ—‘ Removing old chunks: {file_hash}"
    )

    conn = get_conn()

    try:

        with conn.cursor() as cur:

            cur.execute("""
                DELETE FROM upsc_chunks
                WHERE file_hash=%s
            """, (file_hash,))

        conn.commit()

    finally:

        release_conn(conn)

# ==========================================================
# MANIFEST
# ==========================================================

def upsert_manifest(
    file_hash,
    filename,
    chunk_count,
    vector_inserted,
    postgres_inserted
):

    conn = get_conn()

    try:

        with conn.cursor() as cur:

            cur.execute("""
                INSERT INTO ingestion_manifest (
                    file_hash,
                    filename,
                    chunk_count,
                    vector_inserted,
                    postgres_inserted
                )
                VALUES (%s, %s, %s, %s, %s)

                ON CONFLICT (file_hash)

                DO UPDATE SET
                    chunk_count=EXCLUDED.chunk_count,
                    vector_inserted=EXCLUDED.vector_inserted,
                    postgres_inserted=EXCLUDED.postgres_inserted,
                    updated_at=NOW()
            """, (
                file_hash,
                filename,
                chunk_count,
                vector_inserted,
                postgres_inserted
            ))

        conn.commit()

    finally:

        release_conn(conn)

# ==========================================================
# EMBEDDINGS (Gemini API)
# ==========================================================

EMBED_SUB_BATCH_SIZE = int(
    os.getenv("EMBED_SUB_BATCH_SIZE", "20")
)


EMBED_BATCH_DELAY = int(
    os.getenv("EMBED_BATCH_DELAY", "5")
)


def generate_embeddings(chunks):

    if not chunks:
        return np.empty(
            (0, EMBED_DIM),
            dtype=np.float32
        )

    all_embeddings = []
    for i in range(0, len(chunks), EMBED_SUB_BATCH_SIZE):
        if i > 0:
            logger.info(
                f"Waiting {EMBED_BATCH_DELAY}s between embed batches..."
            )
            time.sleep(EMBED_BATCH_DELAY)
        sub_batch = chunks[i : i + EMBED_SUB_BATCH_SIZE]
        batch_embs = _gemini_embed_batch(
            sub_batch, task_type=GEMINI_EMBED_TASK
        )
        all_embeddings.extend(batch_embs)

    return np.asarray(
        all_embeddings,
        dtype=np.float32
    )

def pgvector_literal(embedding):
    values = (
        embedding.tolist()
        if hasattr(embedding, "tolist")
        else list(embedding)
    )

    return "[" + ",".join(str(float(value)) for value in values) + "]"

def insert_postgres_rows(rows):

    def run():

        conn = get_conn()

        try:

            with conn.cursor() as cur:

                cur.execute(
                    "SET LOCAL statement_timeout = %s",
                    (POSTGRES_STATEMENT_TIMEOUT,)
                )

                execute_batch(
                    cur,
                    """
                    INSERT INTO upsc_chunks (
                        id,
                        chunk,
                        topic,
                        difficulty,
                        source_file,
                        file_hash,
                        subject_id,
                        chunk_index,
                        page_number,
                        chunk_version,
                        embedding,
                        search_vector,
                        heading_hierarchy,
                        parent_chunk,
                        is_parent_chunk
                    )
                    VALUES (
                        %s,
                        %s,
                        %s,
                        %s,
                        %s,
                        %s,
                        %s,
                        %s,
                        %s,
                        %s,
                        %s::vector,
                        to_tsvector(
                            'english',
                            %s
                        ),
                        %s::jsonb,
                        %s,
                        %s
                    )

                    ON CONFLICT (id, subject_id)

                    DO UPDATE SET
                        chunk=EXCLUDED.chunk,
                        topic=EXCLUDED.topic,
                        difficulty=EXCLUDED.difficulty,
                        source_file=EXCLUDED.source_file,
                        file_hash=EXCLUDED.file_hash,
                        subject_id=EXCLUDED.subject_id,
                        chunk_index=EXCLUDED.chunk_index,
                        page_number=EXCLUDED.page_number,
                        chunk_version=EXCLUDED.chunk_version,
                        embedding=EXCLUDED.embedding,
                        search_vector=EXCLUDED.search_vector,
                        heading_hierarchy=EXCLUDED.heading_hierarchy,
                        parent_chunk=EXCLUDED.parent_chunk,
                        is_parent_chunk=EXCLUDED.is_parent_chunk
                    """,
                    rows,
                    page_size=BATCH_SIZE
                )

            conn.commit()

        except Exception:

            conn.rollback()

            raise

        finally:

            release_conn(conn)

    retry_operation(
        "PostgreSQL insert",
        run,
        DB_MAX_RETRIES,
        DB_RETRY_DELAY
    )

# ==========================================================
# PROCESS FILE
# ==========================================================

root_folder = None

def process_file(file, subject_id=None):

    global root_folder
    if subject_id:
        subject_id = str(subject_id).strip().lower()
        fname = f"{subject_id}/{file.name}"
    elif root_folder:
        try:
            rel = file.relative_to(root_folder)
            fname = str(rel)
            if len(rel.parts) >= 2:
                subject_id = rel.parts[0].lower()
            else:
                subject_id = root_folder.name
        except ValueError:
            fname = file.name
            subject_id = root_folder.name if root_folder else "general"
    else:
        fname = file.name
        subject_id = "general"

    suffix = file.suffix.lower()

    if suffix not in [
        ".txt",
        ".pdf",
        ".docx"
    ]:
        return

    logger.info(
        f"ðŸ“„ Processing {fname}"
    )

    file_hash = file_checksum(file)

    if is_document_indexed(file_hash):

        logger.info(
            f"⏭️  Skipping {fname} — already indexed"
        )

        return

    update_document_status(
        file_hash,
        fname,
        "processing"
    )

    set_document_subject_id(
        file_hash,
        subject_id
    )

    try:

        # ==================================================
        # READ FILE
        # ==================================================

        page_records = []

        if suffix == ".pdf":

            all_chunks = []
            all_ids = []
            all_metas = []
            all_embeddings = []
            batch_page_records = []
            batch_texts = []
            global_chunk_index = 0

            for page_record in read_pdf_pages(file):

                batch_page_records.append(page_record)
                batch_texts.append(page_record["text"])

                if len(batch_page_records) >= PDF_PAGE_BATCH_SIZE:

                    if ENABLE_VISUAL_VERBALIZATION:
                        enhance_batch_with_visuals(
                            batch_page_records, batch_texts
                        )

                    batch_text = "\n\n".join(batch_texts)
                    batch_chunks = chunk_text(batch_text)

                    if batch_chunks:

                        chunk_texts_batch = [c["chunk_text"] for c in batch_chunks]
                        embeddings_batch = generate_embeddings(chunk_texts_batch)

                        for j, chunk in enumerate(batch_chunks):

                            page_number = infer_chunk_page_number(
                                chunk["chunk_text"],
                                batch_page_records
                            )

                            cid = SHA256.new(
                                f"{file_hash}_{global_chunk_index}".encode()
                            ).hexdigest()

                            all_ids.append(cid)
                            all_chunks.append(chunk)
                            all_embeddings.append(embeddings_batch[j])

                            all_metas.append({
                                "subject_id": subject_id,
                                "topic": detect_topic(chunk["chunk_text"]),
                                "difficulty": detect_difficulty(chunk["chunk_text"]),
                                "source_file": fname,
                                "file_hash": file_hash,
                                "chunk_index": global_chunk_index,
                                "page_number": page_number or 0,
                                "version": 1,
                                "created_at": datetime.utcnow().isoformat(),
                                "heading_hierarchy": chunk.get("heading_hierarchy", []),
                                "parent_text": chunk.get("parent_text", ""),
                                "is_parent_chunk": chunk.get("is_parent_chunk", False),
                            })

                            global_chunk_index += 1

                    batch_page_records.clear()
                    batch_texts.clear()

            if batch_texts:

                if ENABLE_VISUAL_VERBALIZATION:
                    enhance_batch_with_visuals(
                        batch_page_records, batch_texts
                    )

                batch_text = "\n\n".join(batch_texts)
                batch_chunks = chunk_text(batch_text)

                if batch_chunks:

                    chunk_texts_batch = [c["chunk_text"] for c in batch_chunks]
                    embeddings_batch = generate_embeddings(chunk_texts_batch)

                    for j, chunk in enumerate(batch_chunks):

                        page_number = infer_chunk_page_number(
                            chunk["chunk_text"],
                            batch_page_records
                        )

                        cid = SHA256.new(
                            f"{file_hash}_{global_chunk_index}".encode()
                        ).hexdigest()

                        all_ids.append(cid)
                        all_chunks.append(chunk)
                        all_embeddings.append(embeddings_batch[j])

                        all_metas.append({
                            "subject_id": subject_id,
                            "topic": detect_topic(chunk["chunk_text"]),
                            "difficulty": detect_difficulty(chunk["chunk_text"]),
                            "source_file": fname,
                            "file_hash": file_hash,
                            "chunk_index": global_chunk_index,
                            "page_number": page_number or 0,
                            "version": 1,
                            "created_at": datetime.utcnow().isoformat(),
                            "heading_hierarchy": chunk.get("heading_hierarchy", []),
                            "parent_text": chunk.get("parent_text", ""),
                            "is_parent_chunk": chunk.get("is_parent_chunk", False),
                        })

                        global_chunk_index += 1

            if not all_chunks:

                raise Exception(
                    "No valid chunks"
                )

            logger.info(
                f"ðŸ" '{fname} -> {len(all_chunks)} chunks'
            )

            delete_existing_file_data(file_hash)

            rows = []

            for idx, (
                cid,
                chunk,
                meta,
                emb
            ) in enumerate(
                zip(
                    all_ids,
                    all_chunks,
                    all_metas,
                    all_embeddings
                )
            ):

                rows.append((
                    cid,
                    chunk["chunk_text"],
                    meta["topic"],
                    meta["difficulty"],
                    meta["source_file"],
                    file_hash,
                    meta["subject_id"],
                    idx,
                    meta.get("page_number") or None,
                    1,
                    pgvector_literal(emb),
                    chunk["chunk_text"],
                    json.dumps(meta.get("heading_hierarchy", [])),
                    meta.get("parent_text", ""),
                    meta.get("is_parent_chunk", False),
                ))

            insert_postgres_rows(rows)

            upsert_manifest(
                file_hash,
                fname,
                len(all_chunks),
                True,
                True
            )

            update_document_status(
                file_hash,
                fname,
                "indexed"
            )

            logger.info(
                f"âœ… Finished: {fname}"
            )

            del all_chunks
            del all_ids
            del all_metas
            del all_embeddings
            gc.collect()

            return

        elif suffix == ".txt":

            block_generator = read_txt_blocks(file)

        else:

            block_generator = read_docx_blocks(file)

        all_chunks = []
        all_ids = []
        all_metas = []
        all_embeddings = []
        global_chunk_index = 0
        any_block = False

        for block in block_generator:

            any_block = True
            batch_text = block["text"]
            batch_chunks = chunk_text(batch_text)

            if batch_chunks:

                chunk_texts_batch = [c["chunk_text"] for c in batch_chunks]
                embeddings_batch = generate_embeddings(chunk_texts_batch)

                for j, chunk in enumerate(batch_chunks):

                    cid = SHA256.new(
                        f"{file_hash}_{global_chunk_index}".encode()
                    ).hexdigest()

                    all_ids.append(cid)
                    all_chunks.append(chunk)
                    all_embeddings.append(embeddings_batch[j])

                    all_metas.append({
                        "subject_id": subject_id,
                        "topic": detect_topic(chunk["chunk_text"]),
                        "difficulty": detect_difficulty(chunk["chunk_text"]),
                        "source_file": fname,
                        "file_hash": file_hash,
                        "chunk_index": global_chunk_index,
                        "page_number": None,
                        "version": 1,
                        "created_at": datetime.utcnow().isoformat(),
                        "heading_hierarchy": chunk.get("heading_hierarchy", []),
                        "parent_text": chunk.get("parent_text", ""),
                        "is_parent_chunk": chunk.get("is_parent_chunk", False),
                    })

                    global_chunk_index += 1

        if not any_block:

            raise Exception(
                "Empty document"
            )

        if not all_chunks:

            raise Exception(
                "No valid chunks"
            )

        logger.info(
            f"ðŸ" '{fname} -> {len(all_chunks)} chunks'
        )

        delete_existing_file_data(file_hash)

        rows = []

        for idx, (
            cid,
            chunk,
            meta,
            emb
        ) in enumerate(
            zip(
                all_ids,
                all_chunks,
                all_metas,
                all_embeddings
            )
        ):

            rows.append((
                cid,
                chunk["chunk_text"],
                meta["topic"],
                meta["difficulty"],
                meta["source_file"],
                file_hash,
                meta["subject_id"],
                idx,
                None,
                1,
                pgvector_literal(emb),
                chunk["chunk_text"],
                json.dumps(meta.get("heading_hierarchy", [])),
                meta.get("parent_text", ""),
                meta.get("is_parent_chunk", False),
            ))

        insert_postgres_rows(rows)

        upsert_manifest(
            file_hash,
            fname,
            len(all_chunks),
            True,
            True
        )

        update_document_status(
            file_hash,
            fname,
            "indexed"
        )

        logger.info(
            f"âœ… Finished: {fname}"
        )

        del all_chunks
        del all_ids
        del all_metas
        del all_embeddings
        gc.collect()

    except Exception as e:

        logger.exception(
            f"âŒ Failed processing {fname}: {e}"
        )

        update_document_status(
            file_hash,
            fname,
            "failed",
            str(e)
        )

# ==========================================================
# WORKER
# ==========================================================

def ingestion_worker():

    while True:

        file = ingestion_queue.get()

        if file is None:
            break

        try:

            process_file(file)

        except Exception as e:

            logger.exception(
                f"Worker failure: {e}"
            )

        finally:

            ingestion_queue.task_done()

# ==========================================================
# INGESTION
# ==========================================================

def ingest_folder(folder):

    validate_upload_auth()

    ensure_tables()

    global root_folder
    folder_path = Path(folder)
    root_folder = folder_path
    folder_name = folder_path.name

    total_files = 0

    # ======================================================
    # THREAD POOL
    # ======================================================

    workers = []

    for _ in range(MAX_WORKERS):

        worker = threading.Thread(
            target=ingestion_worker,
            daemon=True
        )

        worker.start()

        workers.append(worker)

    # ======================================================
    # QUEUE FILES
    # ======================================================

    for root, dirs, files in os.walk(folder_path):

        for fname in files:

            file = Path(root) / fname

            ingestion_queue.put(file)

            total_files += 1

    logger.info(
        f"ðŸ“¦ Queued files: {total_files}"
    )

    ingestion_queue.join()

    # ======================================================
    # ORPHAN CLEANUP: remove DB records for deleted files
    # ======================================================

    logger.info(
        "ðŸ§¹ Syncing DB state with disk..."
    )

    # Collect current file hashes on disk per subject
    disk_hashes_by_subject = {}

    for root, dirs, files in os.walk(folder_path):

        rel = Path(root).relative_to(folder_path)

        rel_parts = rel.parts

        if len(rel_parts) >= 2:
            subject_id = rel_parts[0].lower()
        else:
            subject_id = folder_path.name

        for fname in files:

            fpath = Path(root) / fname

            suffix = fpath.suffix.lower()

            if suffix not in [
                ".txt",
                ".pdf",
                ".docx"
            ]:
                continue

            fhash = file_checksum(fpath)

            disk_hashes_by_subject.setdefault(
                subject_id,
                set()
            ).add(fhash)

    conn = get_conn()

    try:

        with conn.cursor() as cur:

            subject_keys = list(disk_hashes_by_subject.keys())

            cur.execute(
                "SELECT file_hash, subject_id "
                "FROM documents "
                "WHERE status='indexed' "
                "AND subject_id = ANY(%s)",
                (subject_keys,)
            )

            matched_docs = cur.fetchall()

            for db_hash, db_subj in matched_docs:

                disk_set = disk_hashes_by_subject.get(
                    db_subj,
                    set()
                )

                if db_hash not in disk_set:

                    cur.execute(
                        "SELECT file_name FROM documents "
                        "WHERE file_hash=%s",
                        (db_hash,)
                    )

                    row = cur.fetchone()

                    orphan_name = (
                        row[0] if row else db_hash[:12]
                    )

                    logger.info(
                        f"ðŸ—‘  Orphaned: {orphan_name} "
                        f"in {db_subj} — "
                        "file deleted from disk"
                    )

                    # Remove chunks first, then document
                    cur.execute(
                        "DELETE FROM upsc_chunks "
                        "WHERE file_hash=%s",
                        (db_hash,)
                    )

                    cur.execute(
                        "DELETE FROM documents "
                        "WHERE file_hash=%s",
                        (db_hash,)
                    )

        conn.commit()

    finally:

        release_conn(conn)

    logger.info(
        "ðŸ Ingestion completed"
    )

# ==========================================================
# ENTRY
# ==========================================================

if __name__ == "__main__":

    import argparse

    parser = argparse.ArgumentParser()

    parser.add_argument(
        "--folder",
        required=True
    )

    args = parser.parse_args()

    logger.info(
        f"ðŸš€ Starting ingestion: {args.folder}"
    )

    start = time.time()

    ingest_folder(args.folder)

    elapsed = round(
        time.time() - start,
        2
    )

    logger.info(
        f"âš¡ Total ingestion time: {elapsed}s"
    )
